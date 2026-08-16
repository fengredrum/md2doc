#!/usr/bin/env python3
"""
使用 pandoc 将 Markdown 文件转换为 Word (.docx) 文件，输出符合公文格式规范。

格式规范（来自 格式要求.md）：
    - H1/H2 标题: 黑体 三号(16pt) 加粗 左对齐 固定行距28磅
    - H3 标题: 黑体 四号(14pt) 加粗 左对齐 固定行距28磅
    - 正文: 仿宋_GB2312 四号(14pt) 两端对齐 首行缩进2字符 固定行距28磅
    - 表格: 仿宋_GB2312 五号(10.5pt) 居中对齐 单倍行距

用法:
    uv run python scripts/convert_md2docx.py <input.md> [--output OUTPUT.docx] [--reference REFERENCE.docx]

选项:
    --output, -o            输出路径（默认：与输入同名的 .docx 文件）
    --reference, -r         额外的参考样式模板 .docx（会叠加在内置格式之上）
    --no-format             跳过内置格式规范（退回纯 pandoc 行为）
    --no-cache-template     强制重新生成内置格式模板
    --extract-media DIR     提取媒体文件到指定目录

依赖:
    - pandoc（需在 PATH 中）
    - python-docx + lxml（uv pip install python-docx）
"""

import subprocess
import sys
import os
import re
import html
from html.parser import HTMLParser
from pathlib import Path


# ──────────────────────────────────────────────
# 依赖管理
# ──────────────────────────────────────────────

def _check_pydocx():
    """检查 python-docx 是否可用；不可用则提示安装并退出。"""
    try:
        import docx  # noqa: F401
    except ImportError:
        print("=" * 60)
        print("缺少依赖：python-docx")
        print("请运行以下命令安装：")
        print("  uv pip install python-docx")
        print("=" * 60)
        sys.exit(1)


# ──────────────────────────────────────────────
# HTML 表格解析
# ──────────────────────────────────────────────

class HTMLTableParser(HTMLParser):
    """解析 HTML <table> 元素，提取结构化数据（含 colspan/rowspan）。

    处理 <table>, <tr>, <td>, <th>, <br>, <b> 标签，
    以及 HTML 实体 (&nbsp;, &#160; 等)。
    支持内联 <b> 粗体标记，通过 segments 字段保留粗体/非粗体文本段。
    """

    def __init__(self):
        super().__init__()
        self.tables = []
        self._current_table = None
        self._current_row = None
        self._current_cell = None
        self._in_table = False
        self._in_row = False
        self._in_cell = False
        self._in_bold = False
        # 部分段列表: [{'text': ..., 'bold': bool}, ...]
        self._segments = []

    def _append_text(self, text):
        """将文本追加到当前段。当前没有段时自动创建。"""
        if not self._segments:
            self._segments.append({'text': '', 'bold': self._in_bold})
        # 若粗体状态变化则创建新段
        if self._segments[-1]['bold'] != self._in_bold:
            self._segments.append({'text': '', 'bold': self._in_bold})
        self._segments[-1]['text'] += text

    def _finalize_segments(self):
        """收尾部分段：清理空白、合并相邻同粗体状态段、构建纯文本。"""
        # 1. 清理每个段的文本空白
        for seg in self._segments:
            seg['text'] = re.sub(r'[ \t]+', ' ', seg['text']).strip()
        # 2. 过滤空段
        self._segments = [s for s in self._segments if s['text']]
        # 3. 合并相邻同状态段
        merged = []
        for seg in self._segments:
            if merged and merged[-1]['bold'] == seg['bold']:
                merged[-1]['text'] += seg['text']
            else:
                merged.append(seg)
        self._segments = merged

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        if tag == 'table':
            self._in_table = True
            self._current_table = {'rows': []}
        elif tag == 'tr' and self._in_table:
            self._in_row = True
            self._current_row = []
        elif tag in ('td', 'th') and self._in_row:
            self._in_cell = True
            colspan = int(attrs_dict.get('colspan', 1))
            rowspan = int(attrs_dict.get('rowspan', 1))
            self._segments = []
            self._in_bold = False
            self._current_cell = {
                'text': '',
                'segments': [],
                'colspan': colspan,
                'rowspan': rowspan,
                'is_header': (tag == 'th'),
            }
        elif tag == 'b' and self._in_cell:
            self._in_bold = True
        elif tag == 'br' and self._in_cell:
            self._append_text('\n')

    def handle_endtag(self, tag):
        if tag == 'table' and self._in_table:
            self._in_table = False
            if self._current_table and self._current_table['rows']:
                self.tables.append(self._current_table)
            self._current_table = None
        elif tag == 'tr' and self._in_row:
            self._in_row = False
            if self._current_table is not None and self._current_row is not None:
                self._current_table['rows'].append(self._current_row)
            self._current_row = None
        elif tag in ('td', 'th') and self._in_cell:
            self._in_cell = False
            if self._current_row is not None and self._current_cell is not None:
                self._finalize_segments()
                self._current_cell['segments'] = self._segments
                # 构建兼容纯文本字段（保留向后兼容）
                full_text = ''.join(s['text'] for s in self._segments)
                self._current_cell['text'] = full_text
                self._current_row.append(self._current_cell)
            self._current_cell = None
            self._segments = []
            self._in_bold = False
        elif tag == 'b' and self._in_cell:
            self._in_bold = False

    def handle_data(self, data):
        if self._in_cell and self._current_cell is not None:
            self._append_text(data)

    def handle_entityref(self, name):
        """处理命名实体如 &nbsp; &lt; &gt; &amp;"""
        entity_map = {
            'nbsp': ' ', 'lt': '<', 'gt': '>', 'amp': '&',
            'quot': '"', 'apos': "'",
        }
        if self._in_cell and self._current_cell is not None:
            self._append_text(entity_map.get(name, f'&{name};'))

    def handle_charref(self, name):
        """处理数字字符引用如 &#160; &#x4E2D;"""
        try:
            if name.startswith('x') or name.startswith('X'):
                char = chr(int(name[1:], 16))
            else:
                char = chr(int(name))
        except (ValueError, OverflowError):
            char = f'&#{name};'
        if self._in_cell and self._current_cell is not None:
            self._append_text(char)

    def error(self, message):
        """静默忽略 HTML 解析错误。"""
        pass


def extract_html_tables(md_content):
    """从 Markdown 内容中提取所有 HTML <table> 块，替换为占位符。

    仅成功解析的表格会被替换为占位符；解析失败的表格保留原始 HTML。

    Args:
        md_content: 原始 Markdown 文本

    Returns:
        tuple: (new_content, tables)
        - new_content: 将 HTML 表格替换为占位符后的内容
        - tables: 成功解析的表格数据列表，每项为 {'rows': [[cell_dict, ...], ...]}
          每个 cell_dict 包含: text, colspan, rowspan, is_header
    """
    table_pattern = re.compile(
        r'<table[\s>][\s\S]*?</table>',
        re.IGNORECASE
    )

    matches = list(table_pattern.finditer(md_content))
    if not matches:
        return md_content, []

    tables = []
    placeholder_map = {}  # {start_pos: (end_pos, placeholder)}

    for idx, match in enumerate(matches):
        html_table = match.group(0)
        parser = HTMLTableParser()
        try:
            parser.feed(html_table)
            parsed = parser.tables[0] if parser.tables else None
        except Exception:
            parsed = None

        if parsed is not None:
            tables.append(parsed)
            placeholder = f'\n\n[HTML_TABLE_{len(tables) - 1}]\n\n'
            placeholder_map[match.start()] = (match.end(), placeholder)

    # 无成功解析的表格，直接返回原文
    if not placeholder_map:
        return md_content, []

    # 从左到右构建新内容
    new_parts = []
    last_end = 0
    for start in sorted(placeholder_map.keys()):
        end, placeholder = placeholder_map[start]
        new_parts.append(md_content[last_end:start])
        new_parts.append(placeholder)
        last_end = end
    new_parts.append(md_content[last_end:])
    new_content = ''.join(new_parts)

    return new_content, tables


# ──────────────────────────────────────────────
# pandoc 查找
# ──────────────────────────────────────────────

def find_pandoc():
    """查找 pandoc 可执行文件"""
    local_bin = Path.home() / ".local" / "bin" / "pandoc"
    if local_bin.exists():
        return str(local_bin)
    return "pandoc"


# ──────────────────────────────────────────────
# 参考模板生成
# ──────────────────────────────────────────────

# ── 默认格式规范（当未指定 --format-spec 时的回退方案） ──

DEFAULT_FORMAT_SPEC = {
    'Heading 1': {
        'font_name': '黑体',
        'font_size_pt': 16,      # 三号
        'bold': True,
        'alignment': 'left',
        'line_spacing_rule': 'exactly',
        'line_spacing_pt': 28,
        'space_before_pt': 0,
        'space_after_pt': 0,
        'outline_level': 0,       # 1级
        'first_line_indent_chars': None,
    },
    'Heading 2': {
        'font_name': '黑体',
        'font_size_pt': 16,      # 三号
        'bold': True,
        'alignment': 'left',
        'line_spacing_rule': 'exactly',
        'line_spacing_pt': 28,
        'space_before_pt': 0,
        'space_after_pt': 0,
        'outline_level': 1,       # 2级
        'first_line_indent_chars': None,
    },
    'Heading 3': {
        'font_name': '黑体',
        'font_size_pt': 14,      # 四号
        'bold': True,
        'alignment': 'left',
        'line_spacing_rule': 'exactly',
        'line_spacing_pt': 28,
        'space_before_pt': 0,
        'space_after_pt': 0,
        'outline_level': 2,       # 3级
        'first_line_indent_chars': None,
    },
    'Normal': {
        'font_name': '仿宋_GB2312',
        'font_size_pt': 14,      # 四号
        'bold': False,
        'alignment': 'justify',
        'line_spacing_rule': 'exactly',
        'line_spacing_pt': 28,
        'space_before_pt': 0,
        'space_after_pt': 0,
        'outline_level': None,    # 正文文本（无大纲级别）
        'first_line_indent_chars': '200',  # 2字符 (单位: 1/100字符)
    },
    # pandoc 默认将正文段落映射为 Body Text 而非 Normal
    'Body Text': {
        'font_name': '仿宋_GB2312',
        'font_size_pt': 14,      # 四号
        'bold': False,
        'alignment': 'justify',
        'line_spacing_rule': 'exactly',
        'line_spacing_pt': 28,
        'space_before_pt': 0,
        'space_after_pt': 0,
        'outline_level': None,    # 正文文本（无大纲级别）
        'first_line_indent_chars': '200',  # 2字符 (单位: 1/100字符)
    },
    'Compact': {
        'font_name': '仿宋_GB2312',
        'font_size_pt': 10.5,    # 五号
        'bold': False,
        'alignment': 'center',
        'line_spacing_rule': 'single',
        'line_spacing_pt': None,
        'space_before_pt': 0,
        'space_after_pt': 0,
        'outline_level': None,
        'first_line_indent_chars': None,
    },
}


# ── 格式要求.md 解析 ──

# 中文字号 → 磅值 (pt)
_CN_FONT_SIZE_TO_PT = {
    '三号': 16,
    '小三': 15,
    '四号': 14,
    '五号': 10.5,
    '小四': 12,
    '小五': 9,
}

# 中文对齐方式 → python-docx 标识
_CN_ALIGNMENT_MAP = {
    '左对齐': 'left',
    '两端对齐': 'justify',
    '居中对齐': 'center',
    '右对齐': 'right',
}

# 格式要求.md 节标题前缀 → pandoc 样式名列表
_SECTION_STYLE_MAP = {
    'H1': ['Heading 1'],
    'H2': ['Heading 2'],
    'H3': ['Heading 3'],
    '正文': ['Normal', 'Body Text'],
    '表格': ['Compact'],
}


def _parse_font_field(font_str):
    """解析字体字段，如 "黑体，三号，加粗" 或 "仿宋_GB2312，四号"。

    Returns:
        (font_name, font_size_pt, bold)
    """
    if not font_str:
        return ('仿宋_GB2312', 14, False)

    # 处理 Markdown 转义: "\_" → "_"
    font_str = font_str.replace(r'\_', '_')

    parts = [p.strip() for p in font_str.split('，')]
    font_name = parts[0] if len(parts) >= 1 else '仿宋_GB2312'
    size_cn = parts[1] if len(parts) >= 2 else '四号'
    bold = len(parts) >= 3 and '加粗' in parts[2]

    font_size_pt = _CN_FONT_SIZE_TO_PT.get(size_cn, 14)
    return (font_name, font_size_pt, bold)


def _parse_line_spacing(text):
    """解析行距字段。

    Returns:
        (rule_str, value_pt_or_None)
    """
    if '固定值' in text:
        m = re.search(r'(\d+(?:\.\d+)?)\s*磅', text)
        if m:
            return ('exactly', float(m.group(1)))
        return ('exactly', 28)
    elif '单倍' in text:
        return ('single', None)
    elif '双倍' in text:
        return ('double', None)
    elif '多倍' in text:
        return ('multiple', None)
    return ('exactly', 28)


def _parse_spacing(text):
    """解析间距字段，如 "段前0行，段后0行"。

    Returns:
        (space_before_pt, space_after_pt)
    """
    before_m = re.search(r'段前\s*(\d+(?:\.\d+)?)\s*行', text)
    after_m = re.search(r'段后\s*(\d+(?:\.\d+)?)\s*行', text)
    before = float(before_m.group(1)) if before_m else 0.0
    after = float(after_m.group(1)) if after_m else 0.0
    return (before, after)


def _parse_outline_level(text):
    """解析大纲级别字段。

    "1级" → 0, "2级" → 1, "正文文本" → None, "无" → None
    """
    if '正文文本' in text or text == '无' or not text:
        return None
    m = re.search(r'(\d+)\s*级', text)
    if m:
        return int(m.group(1)) - 1  # 1-based → 0-based
    return None


def _parse_special_format(text):
    """解析特殊格式字段。

    "首行缩进2字符" → "200" (2 × 100, 单位 1/100 字符)
    "无" → None
    """
    if not text or text == '无':
        return None
    m = re.search(r'首行缩进\s*(\d+(?:\.\d+)?)\s*字符', text)
    if m:
        return str(int(float(m.group(1)) * 100))
    return None


def parse_format_md(filepath):
    """解析 格式要求.md，返回 FORMAT_SPEC 格式的字典。

    Args:
        filepath: 格式要求.md 的路径

    Returns:
        dict: 与 DEFAULT_FORMAT_SPEC 结构一致的样式配置字典

    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 文件内容无法解析
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    # 按 "## " 分节
    sections = re.split(r'\n(?=## )', content)
    format_spec = {}

    for section in sections:
        lines = section.strip().split('\n')
        if not lines or not lines[0].startswith('## '):
            continue

        section_name = lines[0][3:].strip()  # 去掉 "## " 前缀

        # 匹配节标题到样式名
        matched_key = None
        for map_key in _SECTION_STYLE_MAP:
            if section_name.startswith(map_key):
                matched_key = map_key
                break

        if matched_key is None:
            print(f"警告: 未知的格式章节 '{section_name}'，已跳过")
            continue

        style_names = _SECTION_STYLE_MAP[matched_key]

        # 解析属性行
        props = {}
        for line in lines[1:]:
            line = line.strip()
            if not line.startswith('- '):
                continue
            content_line = line[2:]
            if '：' in content_line:
                key, value = content_line.split('：', 1)
                props[key.strip()] = value.strip()

        # 构建配置
        font_name, size_pt, bold = _parse_font_field(props.get('字体', ''))
        alignment = _CN_ALIGNMENT_MAP.get(props.get('对齐方式', ''), 'left')
        line_rule, line_pt = _parse_line_spacing(props.get('行距', ''))
        outline = _parse_outline_level(props.get('大纲级别', ''))
        indent = _parse_special_format(props.get('特殊格式', ''))
        before, after = _parse_spacing(props.get('间距', ''))

        config = {
            'font_name': font_name,
            'font_size_pt': size_pt,
            'bold': bold,
            'alignment': alignment,
            'line_spacing_rule': line_rule,
            'line_spacing_pt': line_pt,
            'space_before_pt': before,
            'space_after_pt': after,
            'outline_level': outline,
            'first_line_indent_chars': indent,
        }

        for sn in style_names:
            format_spec[sn] = config

    if not format_spec:
        raise ValueError("格式要求.md 未生成任何样式配置，请检查文件内容")

    return format_spec


def get_format_spec(format_md_path=None):
    """获取格式规范配置。

    若提供 format_md_path 则从文件动态解析；
    否则回退到内置 DEFAULT_FORMAT_SPEC。

    Args:
        format_md_path: 格式要求.md 的路径（可选）

    Returns:
        dict: 样式配置字典
    """
    if format_md_path is None:
        return DEFAULT_FORMAT_SPEC

    if not os.path.exists(format_md_path):
        print(f"警告: 格式规范文件不存在: {format_md_path}")
        print("将使用内置默认格式规范 (DEFAULT_FORMAT_SPEC)")
        return DEFAULT_FORMAT_SPEC

    try:
        spec = parse_format_md(format_md_path)
        print(f"已从 {format_md_path} 加载格式规范")
        return spec
    except Exception as e:
        print(f"警告: 解析格式规范文件失败: {e}")
        print("将使用内置默认格式规范 (DEFAULT_FORMAT_SPEC)")
        return DEFAULT_FORMAT_SPEC


def _get_alignment_enum(alignment_str):
    """将字符串对齐方式转为 python-docx 枚举值。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    mapping = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)


def _get_line_spacing_rule_enum(rule_str):
    """将行距规则字符串转为 python-docx 枚举值。"""
    from docx.enum.text import WD_LINE_SPACING
    mapping = {
        'exactly': WD_LINE_SPACING.EXACTLY,
        'single': WD_LINE_SPACING.SINGLE,
        'double': WD_LINE_SPACING.DOUBLE,
        'multiple': WD_LINE_SPACING.MULTIPLE,
    }
    return mapping.get(rule_str, WD_LINE_SPACING.SINGLE)


def _set_east_asian_font(style_or_run_element, font_name):
    """在样式或 run 元素的 rPr 中设置东亚字体 (w:eastAsia)。"""
    from docx.oxml.ns import qn
    from lxml import etree

    # 找到或创建 w:rPr
    rPr = style_or_run_element.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(style_or_run_element, qn('w:rPr'))

    # 找到或创建 w:rFonts
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        # rFonts 应放在 rPr 的第一个位置（在 sz 等之前）
        rFonts = etree.Element(qn('w:rFonts'))
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), font_name)


def _set_first_line_indent_chars(pPr_element, chars_value):
    """在段落属性中设置首行缩进（字符单位）。chars_value 为字符串如 '200'。"""
    from docx.oxml.ns import qn
    from lxml import etree

    ind = pPr_element.find(qn('w:ind'))
    if ind is None:
        ind = etree.SubElement(pPr_element, qn('w:ind'))
    ind.set(qn('w:firstLineChars'), chars_value)


def _clear_first_line_indent(pPr_element):
    """显式清除首行缩进，防止从基样式（Normal）继承缩进。

    Word 样式继承规则：若当前样式的 w:pPr 中没有 w:ind/@w:firstLine
    或 @w:firstLineChars，则会继承 basedOn 样式（如 Normal）的缩进设置。
    因此对于标题等不需要缩进的样式，必须显式写入 w:firstLine="0"。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    ind = pPr_element.find(qn('w:ind'))
    if ind is None:
        ind = etree.SubElement(pPr_element, qn('w:ind'))
    # 清除 twip 单位缩进（可能被其他代码设置）
    if qn('w:firstLine') in ind.attrib:
        del ind.attrib[qn('w:firstLine')]
    # 显式清零字符单位缩进——必须用与 Normal 相同的属性类型 (firstLineChars)
    # 否则 Word 不会正确覆盖基样式的 firstLineChars 继承
    ind.set(qn('w:firstLineChars'), '0')


def _set_outline_level(pPr_element, level):
    """设置大纲级别。0=1级, 1=2级, ..."""
    from docx.oxml.ns import qn
    from lxml import etree

    # 移除现存的 outlineLvl
    for existing in pPr_element.findall(qn('w:outlineLvl')):
        pPr_element.remove(existing)
    outline_lvl = etree.SubElement(pPr_element, qn('w:outlineLvl'))
    outline_lvl.set(qn('w:val'), str(level))


def _apply_format_spec_to_style(style, config):
    """将一份格式配置应用到 python-docx Style 对象上。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    # 字体
    style.font.name = config['font_name']
    style.font.size = Pt(config['font_size_pt'])
    style.font.bold = config['bold']

    # 东亚字体 (通过 XML 操作)
    _set_east_asian_font(style.element, config['font_name'])

    # 字体颜色：显式设为自动（黑色），覆盖 python-docx 默认模板的蓝色标题
    style.font.color.rgb = None  # None = 自动颜色
    # 同时从 XML 中移除 w:color 元素（python-docx 默认模板会在标题样式中预设蓝色）
    rPr = style.element.find(qn('w:rPr'))
    if rPr is not None:
        color_el = rPr.find(qn('w:color'))
        if color_el is not None:
            rPr.remove(color_el)

    # 段落格式
    pf = style.paragraph_format
    pf.alignment = _get_alignment_enum(config['alignment'])
    pf.line_spacing_rule = _get_line_spacing_rule_enum(config['line_spacing_rule'])
    if config['line_spacing_pt'] is not None:
        pf.line_spacing = Pt(config['line_spacing_pt'])
    pf.space_before = Pt(config['space_before_pt'])
    pf.space_after = Pt(config['space_after_pt'])

    # 首行缩进
    # - 有值: 显式设置（正文首行缩进2字符）
    # - None: 显式清零，防止从基样式 Normal 继承缩进（标题）
    pPr = style.element.find(qn('w:pPr'))
    if pPr is not None:
        if config['first_line_indent_chars'] is not None:
            _set_first_line_indent_chars(pPr, config['first_line_indent_chars'])
        else:
            _clear_first_line_indent(pPr)

    # 大纲级别
    if config['outline_level'] is not None:
        pPr = style.element.find(qn('w:pPr'))
        if pPr is not None:
            _set_outline_level(pPr, config['outline_level'])


def _ensure_style_exists(doc, style_name, base_style='Normal'):
    """确保文档中存在某个样式；不存在则基于 base_style 创建。"""
    from docx.oxml.ns import qn

    try:
        return doc.styles[style_name]
    except KeyError:
        # Compact 等样式可能不在默认模板中，需要手动添加
        style = doc.styles.add_style(style_name, 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        style.base_style = doc.styles[base_style]
        return style


def create_reference_docx(output_path, format_spec=None):
    """
    生成符合 格式要求.md 规范的 pandoc 参考模板 .docx。

    pandoc 会将 MD 元素映射到模板中的样式名：
        H1 → Heading 1, H2 → Heading 2, H3 → Heading 3,
        正文 → Normal, 表格文本 → Compact

    Args:
        output_path: 输出 .docx 路径
        format_spec: 格式规范字典（由 get_format_spec() 返回）
    """
    _check_pydocx()
    from docx import Document

    if format_spec is None:
        format_spec = DEFAULT_FORMAT_SPEC

    doc = Document()

    for style_name, config in format_spec.items():
        _ensure_style_exists(doc, style_name)
        style = doc.styles[style_name]
        _apply_format_spec_to_style(style, config)

    # pandoc 要求参考模板至少包含一个段落才能正确读取样式
    doc.add_paragraph(' ', style='Normal')

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ──────────────────────────────────────────────
# HTML 表格 → Word 表格转换
# ──────────────────────────────────────────────

def _set_cell_vertical_alignment(cell, align='center'):
    """设置单元格内容的垂直对齐方式。

    Args:
        cell: python-docx Cell 对象
        align: 'center' (居中), 'top' (顶端), 或 'bottom' (底端)
    """
    from docx.oxml.ns import qn
    from lxml import etree

    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn('w:tcPr'))
        tc.insert(0, tcPr)

    # 移除现存的 vAlign
    for existing in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(existing)

    vAlign = etree.SubElement(tcPr, qn('w:vAlign'))
    vAlign.set(qn('w:val'), align)


def _set_table_autofit_window(table):
    """设置表格宽度为 100% 页宽，根据窗口和内容自动调整列宽。

    移除固定列宽和单元格宽度约束，让 Word 根据窗口大小和单元格内容
    自动计算最佳列宽。对应 Word 中的「根据窗口自动调整表格」行为。

    Args:
        table: python-docx Table 对象
    """
    from docx.oxml.ns import qn
    from lxml import etree

    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, qn('w:tblPr'))
        tbl.insert(0, tblPr)

    # 设置表格宽度为 100%（5000 = 100%）
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')

    # 移除 autofitToContents（如果存在），允许窗口自适应
    autofit = tblPr.find(qn('w:autofitToContents'))
    if autofit is not None:
        tblPr.remove(autofit)

    # 移除 tblGrid 中的固定列宽，让 Word 自动计算
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gridCol in tblGrid.findall(qn('w:gridCol')):
            # 清除固定宽度但保留列定义
            if qn('w:w') in gridCol.attrib:
                del gridCol.attrib[qn('w:w')]

    # 移除所有单元格的固定宽度 (w:tcW)，让 Word 根据内容自动分配列宽
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)


def _create_word_table(doc, table_data, format_spec=None):
    """从解析的 HTML 表格数据创建 python-docx Table 对象。

    处理 colspan/rowspan 单元格合并，<br> 换行，<th> 加粗，
    并应用 Compact 样式：仿宋_GB2312 五号(10.5pt) 居中对齐。

    Args:
        doc: python-docx Document 对象
        table_data: {'rows': [[cell_dict, ...], ...]}
        format_spec: 格式规范字典

    Returns:
        docx.table.Table 或 None
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn

    rows_data = table_data.get('rows', [])
    if not rows_data:
        return None

    # 计算总列数（取所有行中 colspan 之和的最大值）
    num_cols = 0
    for row in rows_data:
        total = sum(cell.get('colspan', 1) for cell in row)
        num_cols = max(num_cols, total)

    if num_cols == 0:
        return None

    num_rows = len(rows_data)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # 设置表格适应窗口宽度
    _set_table_autofit_window(table)

    # 追踪被 rowspan 覆盖的单元格位置
    covered = set()

    # 获取表格字体配置
    table_font = '仿宋_GB2312'
    table_size_pt = 10.5  # 五号
    if format_spec:
        compact = format_spec.get('Compact', {})
        table_font = compact.get('font_name', table_font)
        table_size_pt = compact.get('font_size_pt', table_size_pt)

    for row_idx, row_data in enumerate(rows_data):
        col_idx = 0
        for cell_data in row_data:
            # 跳过被 rowspan 覆盖的列
            while (row_idx, col_idx) in covered:
                col_idx += 1

            if col_idx >= num_cols:
                break

            cell = table.cell(row_idx, col_idx)
            colspan = cell_data.get('colspan', 1)
            rowspan = cell_data.get('rowspan', 1)
            cell_text = cell_data.get('text', '')
            is_header = cell_data.get('is_header', False)

            # 设置单元格垂直居中
            _set_cell_vertical_alignment(cell, 'center')

            # 清除默认空段落，写入格式化文本
            first_para = cell.paragraphs[0]
            first_para.clear()
            first_para.alignment = _get_alignment_enum('center')

            # 清除首行缩进（防止从 Normal 样式继承 2 字符缩进）
            pPr = first_para._element.find(qn('w:pPr'))
            if pPr is not None:
                _clear_first_line_indent(pPr)
            else:
                from lxml import etree as _etree
                pPr = _etree.SubElement(first_para._element, qn('w:pPr'))
                first_para._element.insert(0, pPr)
                _clear_first_line_indent(pPr)

            if cell_text:
                segments = cell_data.get('segments', [])
                if segments:
                    # 使用 segments 保留内联粗体标记
                    for seg in segments:
                        lines = seg['text'].split('\n')
                        for i, line in enumerate(lines):
                            clean_line = re.sub(r'\s+', ' ', line).strip()
                            if not clean_line:
                                continue
                            run = first_para.add_run(clean_line)
                            run.font.name = table_font
                            run.font.size = Pt(table_size_pt)
                            if seg['bold'] or is_header:
                                run.bold = True
                            _set_east_asian_font(run._element, table_font)
                            if i < len(lines) - 1:
                                run.add_break()
                else:
                    # 向后兼容：无 segments 时按行分割
                    lines = cell_text.split('\n')
                    for i, line in enumerate(lines):
                        clean_line = re.sub(r'\s+', ' ', line).strip()
                        run = first_para.add_run(clean_line if clean_line else ' ')
                        run.font.name = table_font
                        run.font.size = Pt(table_size_pt)
                        if is_header:
                            run.bold = True
                        _set_east_asian_font(run._element, table_font)
                        if i < len(lines) - 1:
                            run.add_break()
            else:
                # 空单元格：需要一个空 run 满足 Word 最小内容要求
                run = first_para.add_run(' ')
                run.font.name = table_font
                run.font.size = Pt(table_size_pt)
                if is_header:
                    run.bold = True
                _set_east_asian_font(run._element, table_font)

            # 处理单元格合并（colspan + rowspan 矩形合并）
            if colspan > 1 or rowspan > 1:
                end_row = row_idx + rowspan - 1
                end_col = col_idx + colspan - 1

                # 标记被覆盖的单元格
                for r in range(row_idx, end_row + 1):
                    for c in range(col_idx, end_col + 1):
                        if r != row_idx or c != col_idx:
                            covered.add((r, c))

                # 合并单元格（矩形合并）
                if end_row != row_idx or end_col != col_idx:
                    try:
                        cell.merge(table.cell(end_row, end_col))
                    except Exception:
                        pass

            col_idx += colspan

    return table


def _insert_html_tables(docx_path, tables, format_spec=None):
    """在 DOCX 中查找 HTML 表格占位符 [HTML_TABLE_N] 并替换为 Word 原生表格。

    无 HTML 表格时快速跳过（零开销）。

    Args:
        docx_path: DOCX 文件路径
        tables: 从 extract_html_tables() 返回的成功解析的表格数据列表
        format_spec: 格式规范字典

    Returns:
        int: 成功插入的表格数量
    """
    if not tables:
        return 0

    _check_pydocx()
    from docx import Document

    doc = Document(docx_path)
    body = doc.element.body

    PLACEHOLDER_RE = re.compile(r'^\s*\[HTML_TABLE_(\d+)\]\s*$')

    placeholder_paras = []  # [(para_element, table_index)]

    for para in doc.paragraphs:
        text = para.text.strip()
        m = PLACEHOLDER_RE.match(text)
        if not m:
            continue

        idx = int(m.group(1))
        if idx >= len(tables) or tables[idx] is None:
            continue

        table_data = tables[idx]
        if not table_data.get('rows'):
            continue

        word_table = _create_word_table(doc, table_data, format_spec)
        if word_table is None:
            continue

        tbl_element = word_table._tbl

        # 从文档末尾移除（add_table 将其追加到 body 末尾）
        body.remove(tbl_element)

        # 插入到占位符段落之后
        para_element = para._element
        para_index = list(body).index(para_element)
        body.insert(para_index + 1, tbl_element)

        placeholder_paras.append(para_element)

    # 删除占位符段落
    for elem in placeholder_paras:
        body.remove(elem)

    if placeholder_paras:
        doc.save(docx_path)
        print(f"HTML表格处理: 已插入 {len(placeholder_paras)} 个表格到 Word 文档")

    return len(placeholder_paras)


# ──────────────────────────────────────────────
# 后处理
# ──────────────────────────────────────────────

def _fix_image_paragraph_spacing(doc, format_spec=None):
    """修正包含图片的段落行距，防止图片被裁剪。

    pandoc 将 Markdown 图片转换为 Word 内联图片 (wp:inline)，
    放置在 Normal 或 Body Text 样式段落中。这些段落从样式继承
    "固定值 28 磅" 行距，而图片通常远大于 28 磅，会被 Word 裁剪
    至一行高度，导致图片与后续文字重叠。

    本函数查找所有包含 w:drawing 元素的段落，将其行距显式设为
    单倍行距，使行高能自动扩展以完整显示图片。

    Args:
        doc: python-docx Document 对象
        format_spec: 格式规范字典（保留参数，为未来扩展预留）
    """
    from docx.oxml.ns import qn
    from lxml import etree

    for paragraph in doc.paragraphs:
        # 检查段落是否包含图片（w:drawing 元素）
        drawings = paragraph._element.findall('.//' + qn('w:drawing'))
        if not drawings:
            continue

        # 获取或创建 w:pPr
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = etree.SubElement(paragraph._element, qn('w:pPr'))
            paragraph._element.insert(0, pPr)

        # 获取或创建 w:spacing 元素
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = etree.SubElement(pPr, qn('w:spacing'))

        # 设置为单倍行距（lineRule="auto" 或直接移除固定行距约束）
        # 使用 lineRule="auto" + line="240" (单倍行距的标准值)
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        # 清除可能存在的段前段后约束，保持图片上下间距自然
        # 保留 before/after 让 Word 自行排版


def _get_content_elements(doc):
    """遍历 body 元素，返回内容承载元素列表（跳过纯空白段落）。

    按文档顺序遍历 w:body 的直接子元素（w:p 和 w:tbl），
    过滤掉仅包含空白字符的段落，保留表格和含文本的段落。

    Args:
        doc: python-docx Document 对象

    Returns:
        [{'type': 'paragraph', 'para': Paragraph, 'element': lxml Element}, ...]
    """
    from docx.oxml.ns import qn

    # 构建 python-docx 对象到 lxml 元素的映射
    para_by_elem = {}
    for p in doc.paragraphs:
        para_by_elem[id(p._element)] = p

    table_by_elem = {}
    for t in doc.tables:
        table_by_elem[id(t._tbl)] = t

    elements = []
    body = doc.element.body

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = para_by_elem.get(id(child))
            if para is None:
                continue
            # 跳过纯空白段落（用于模版中作垂直间距的空白段落）
            text = para.text.strip() if para.text else ''
            has_content = bool(text) or any(
                r.text and r.text.strip() for r in para.runs
            )
            if has_content:
                elements.append({'type': 'paragraph', 'para': para, 'element': child})
        elif tag == 'tbl':
            table = table_by_elem.get(id(child))
            if table is not None:
                elements.append({'type': 'table', 'table': table, 'element': child})

    return elements


def _copy_paragraph_format(tpl_para, out_para):
    """从模版段落复制格式到输出段落。

    复制段落级属性（对齐、缩进、间距）和 run 级属性（字体名称、大小），
    但保留输出中已有的粗体/斜体标记，以维护 Markdown 语义格式。

    Args:
        tpl_para: 模版 python-docx Paragraph 对象
        out_para: 输出 python-docx Paragraph 对象
    """
    from docx.oxml.ns import qn
    from lxml import etree

    # ── 段落级属性 ──
    tpl_pPr = tpl_para._element.find(qn('w:pPr'))
    out_pPr = out_para._element.find(qn('w:pPr'))

    if tpl_pPr is not None and out_pPr is not None:
        # 段落对齐 w:jc
        tpl_jc = tpl_pPr.find(qn('w:jc'))
        if tpl_jc is not None:
            out_jc = out_pPr.find(qn('w:jc'))
            if out_jc is None:
                out_jc = etree.SubElement(out_pPr, qn('w:jc'))
            out_jc.set(qn('w:val'), tpl_jc.get(qn('w:val')))

        # 段落缩进 w:ind
        tpl_ind = tpl_pPr.find(qn('w:ind'))
        if tpl_ind is not None:
            out_ind = out_pPr.find(qn('w:ind'))
            if out_ind is None:
                out_ind = etree.SubElement(out_pPr, qn('w:ind'))
            for attr_name, attr_val in tpl_ind.attrib.items():
                out_ind.set(attr_name, attr_val)

        # 段落间距 w:spacing
        tpl_spacing = tpl_pPr.find(qn('w:spacing'))
        if tpl_spacing is not None:
            out_spacing = out_pPr.find(qn('w:spacing'))
            if out_spacing is None:
                out_spacing = etree.SubElement(out_pPr, qn('w:spacing'))
            for attr_name, attr_val in tpl_spacing.attrib.items():
                out_spacing.set(attr_name, attr_val)

    # ── Run 级属性（按位置匹配：输出 run N 应用模版 run N 的格式） ──
    # 若输出 run 多于模版 run，多余 run 保留现有格式；
    # 若模版 run 多于输出 run，多余模版 run 忽略。
    if not out_para.runs:
        return

    for out_idx, out_run in enumerate(out_para.runs):
        if out_idx >= len(tpl_para.runs):
            break

        tpl_run = tpl_para.runs[out_idx]
        tpl_rPr = tpl_run._element.find(qn('w:rPr'))

        # 从模版 run 提取字体/大小信息
        tpl_rFonts = tpl_rPr.find(qn('w:rFonts')) if tpl_rPr is not None else None
        tpl_sz = tpl_rPr.find(qn('w:sz')) if tpl_rPr is not None else None
        tpl_szCs = tpl_rPr.find(qn('w:szCs')) if tpl_rPr is not None else None

        out_rPr = out_run._element.find(qn('w:rPr'))
        if out_rPr is None:
            out_rPr = etree.SubElement(out_run._element, qn('w:rPr'))
            out_run._element.insert(0, out_rPr)

        # 保存输出 run 的粗体标记（Markdown 语义，模版 run 的粗体不覆盖输出）
        saved_b = out_rPr.find(qn('w:b'))
        saved_bCs = out_rPr.find(qn('w:bCs'))
        had_bold = (
            (saved_b is not None and saved_b.get(qn('w:val')) != '0')
            or (saved_bCs is not None and saved_bCs.get(qn('w:val')) != '0')
        )

        # 复制/清除字体名称：模版有则复制，模版无则清除输出以继承样式
        if tpl_rFonts is not None:
            out_rFonts = out_rPr.find(qn('w:rFonts'))
            if out_rFonts is None:
                out_rFonts = etree.SubElement(out_rPr, qn('w:rFonts'))
                out_rPr.insert(0, out_rFonts)
            for attr in ('eastAsia', 'ascii', 'hAnsi', 'cs'):
                val = tpl_rFonts.get(qn(f'w:{attr}'))
                if val:
                    out_rFonts.set(qn(f'w:{attr}'), val)
        else:
            out_rFonts = out_rPr.find(qn('w:rFonts'))
            if out_rFonts is not None:
                out_rPr.remove(out_rFonts)

        # 复制/清除字号：模版有则复制，模版无则清除输出以继承样式
        if tpl_sz is not None:
            out_sz = out_rPr.find(qn('w:sz'))
            if out_sz is None:
                out_sz = etree.SubElement(out_rPr, qn('w:sz'))
            out_sz.set(qn('w:val'), tpl_sz.get(qn('w:val')))
        else:
            out_sz = out_rPr.find(qn('w:sz'))
            if out_sz is not None:
                out_rPr.remove(out_sz)

        if tpl_szCs is not None:
            out_szCs = out_rPr.find(qn('w:szCs'))
            if out_szCs is None:
                out_szCs = etree.SubElement(out_rPr, qn('w:szCs'))
            out_szCs.set(qn('w:val'), tpl_szCs.get(qn('w:val')))
        else:
            out_szCs = out_rPr.find(qn('w:szCs'))
            if out_szCs is not None:
                out_rPr.remove(out_szCs)

        # 恢复输出原有的粗体标记
        if had_bold and saved_b is None:
            etree.SubElement(out_rPr, qn('w:b'))
        # else: 粗体标记已存在或不需要，不做改动


def _copy_cell_width(tpl_cell, out_cell):
    """从模版单元格复制宽度到输出单元格。

    Args:
        tpl_cell: 模版 python-docx Cell 对象
        out_cell: 输出 python-docx Cell 对象
    """
    from docx.oxml.ns import qn
    from lxml import etree

    tpl_tcPr = tpl_cell._tc.find(qn('w:tcPr'))
    out_tcPr = out_cell._tc.find(qn('w:tcPr'))

    if tpl_tcPr is None or out_tcPr is None:
        return

    tpl_tcW = tpl_tcPr.find(qn('w:tcW'))
    if tpl_tcW is not None:
        out_tcW = out_tcPr.find(qn('w:tcW'))
        if out_tcW is None:
            out_tcW = etree.SubElement(out_tcPr, qn('w:tcW'))
        for attr_name, attr_val in tpl_tcW.attrib.items():
            out_tcW.set(attr_name, attr_val)


def _copy_table_format(tpl_table, out_table):
    """从模版表格复制格式到输出表格。

    复制表格样式、边框、列宽、布局设置，
    以及每个单元格的段落/run 格式。

    Args:
        tpl_table: 模版 python-docx Table 对象
        out_table: 输出 python-docx Table 对象
    """
    from docx.oxml.ns import qn
    from lxml import etree
    from copy import deepcopy

    tpl_tbl = tpl_table._tbl
    out_tbl = out_table._tbl

    # ── 表格级属性 ──
    tpl_tblPr = tpl_tbl.find(qn('w:tblPr'))
    out_tblPr = out_tbl.find(qn('w:tblPr'))

    if tpl_tblPr is not None and out_tblPr is not None:
        # 表格样式 w:tblStyle
        tpl_tblStyle = tpl_tblPr.find(qn('w:tblStyle'))
        if tpl_tblStyle is not None:
            out_tblStyle = out_tblPr.find(qn('w:tblStyle'))
            if out_tblStyle is None:
                out_tblStyle = etree.SubElement(out_tblPr, qn('w:tblStyle'))
                out_tblPr.insert(0, out_tblStyle)
            out_tblStyle.set(qn('w:val'), tpl_tblStyle.get(qn('w:val')))

        # 表格宽度 w:tblW
        tpl_tblW = tpl_tblPr.find(qn('w:tblW'))
        if tpl_tblW is not None:
            out_tblW = out_tblPr.find(qn('w:tblW'))
            if out_tblW is None:
                out_tblW = etree.SubElement(out_tblPr, qn('w:tblW'))
            for attr_name, attr_val in tpl_tblW.attrib.items():
                out_tblW.set(attr_name, attr_val)

        # 表格对齐 w:jc（OOXML 中表格居中对齐元素位于 w:tblPr/w:jc）
        tpl_jc = tpl_tblPr.find(qn('w:jc'))
        if tpl_jc is not None:
            out_jc = out_tblPr.find(qn('w:jc'))
            if out_jc is None:
                out_jc = etree.Element(qn('w:jc'))
                # 按 Word schema 顺序插入：紧跟 w:tblW 之后
                anchor = out_tblPr.find(qn('w:tblW'))
                if anchor is not None:
                    anchor.addnext(out_jc)
                else:
                    out_tblPr.insert(0, out_jc)
            out_jc.set(qn('w:val'), tpl_jc.get(qn('w:val')))

        # 表格布局 w:tblLayout
        tpl_tblLayout = tpl_tblPr.find(qn('w:tblLayout'))
        if tpl_tblLayout is not None:
            out_tblLayout = out_tblPr.find(qn('w:tblLayout'))
            if out_tblLayout is None:
                out_tblLayout = etree.SubElement(out_tblPr, qn('w:tblLayout'))
            for attr_name, attr_val in tpl_tblLayout.attrib.items():
                out_tblLayout.set(attr_name, attr_val)

        # 表格边框 w:tblBorders（深拷贝整个边框元素）
        tpl_tblBorders = tpl_tblPr.find(qn('w:tblBorders'))
        if tpl_tblBorders is not None:
            out_tblBorders = out_tblPr.find(qn('w:tblBorders'))
            if out_tblBorders is not None:
                out_tblPr.remove(out_tblBorders)
            out_tblPr.append(deepcopy(tpl_tblBorders))

    # ── 列宽：从模版 tblGrid 复制 ──
    tpl_tblGrid = tpl_tbl.find(qn('w:tblGrid'))
    out_tblGrid = out_tbl.find(qn('w:tblGrid'))
    if tpl_tblGrid is not None and out_tblGrid is not None:
        tpl_cols = tpl_tblGrid.findall(qn('w:gridCol'))
        out_cols = out_tblGrid.findall(qn('w:gridCol'))
        for i, tpl_col in enumerate(tpl_cols):
            if i < len(out_cols):
                w = tpl_col.get(qn('w:w'))
                if w is not None:
                    out_cols[i].set(qn('w:w'), w)

    # ── 单元格级格式 ──
    for row_idx in range(min(len(tpl_table.rows), len(out_table.rows))):
        tpl_row = tpl_table.rows[row_idx]
        out_row = out_table.rows[row_idx]

        for col_idx in range(min(len(tpl_row.cells), len(out_row.cells))):
            tpl_cell = tpl_row.cells[col_idx]
            out_cell = out_row.cells[col_idx]

            # 复制单元格宽度
            _copy_cell_width(tpl_cell, out_cell)

            # 设置垂直居中
            _set_cell_vertical_alignment(out_cell, 'center')

            # 复制每个段落格式：跳过模版单元格中的空白(间距)段落，
            # 按"内容段落"位置匹配（输出 HTML 表格单元格恒为单段落）
            tpl_paras = [p for p in tpl_cell.paragraphs if p.text.strip()]
            out_paras = out_cell.paragraphs
            for p_idx in range(min(len(tpl_paras), len(out_paras))):
                _copy_paragraph_format(tpl_paras[p_idx], out_paras[p_idx])


def _copy_template_formatting(output_path, template_path,
                              copy_paragraphs=True, copy_tables=True):
    """从模版 docx 复制格式到输出 docx 的对应元素。

    按 body 元素顺序进行位置匹配——模版中的空白间距段落会被跳过，
    因此输出元素 N 会匹配到模版中第 N 个内容承载元素。

    Args:
        output_path: pandoc 输出的 .docx 文件路径（原地修改）
        template_path: 模版 .docx 文件路径
        copy_paragraphs: 是否复制段落格式
        copy_tables: 是否复制表格格式
    """
    _check_pydocx()
    from docx import Document

    out_doc = Document(output_path)
    tpl_doc = Document(template_path)

    tpl_elements = _get_content_elements(tpl_doc)
    out_elements = _get_content_elements(out_doc)

    for i, out_elem in enumerate(out_elements):
        if i >= len(tpl_elements):
            break
        tpl_elem = tpl_elements[i]

        if (out_elem['type'] == 'paragraph' and tpl_elem['type'] == 'paragraph'
                and copy_paragraphs):
            _copy_paragraph_format(tpl_elem['para'], out_elem['para'])
        elif (out_elem['type'] == 'table' and tpl_elem['type'] == 'table'
                and copy_tables):
            _copy_table_format(tpl_elem['table'], out_elem['table'])

    out_doc.save(output_path)


def post_process_docx(output_path, format_spec=None, reference_docx_path=None):
    """
    对 pandoc 生成的 DOCX 进行轻量后处理，修正东亚字体设置与图片排版。

    pandoc 在使用参考模板时，会在部分文本运行上写入直接的字体设置，
    覆盖模板中的东亚字体。本函数遍历全文，将跑偏的字体修正回来。

    同时修正图片段落：pandoc 生成的内联图片位于 Normal 样式段落中，
    该样式设置了固定值 28 磅行距，会导致大尺寸图片被裁剪至 28 磅高度，
    造成图片被文字遮挡。本函数将图片段落的行距改为单倍行距以完整显示图片。

    当指定 reference_docx_path 时，会从模版文档复制格式到输出文档，
    并跳过表格单元格的强制格式化，以保持模版中的个性化表格样式。

    Args:
        output_path: 要后处理的 .docx 文件路径
        format_spec: 格式规范字典（由 get_format_spec() 返回）
        reference_docx_path: 模版 .docx 文件路径。指定后从模版复制格式，
            并跳过表格单元格的强制格式化。
    """
    _check_pydocx()
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from lxml import etree

    if format_spec is None:
        format_spec = DEFAULT_FORMAT_SPEC

    doc = Document(output_path)

    # 正文字体（用于表格单元格和未匹配段落）
    DEFAULT_FONT = '仿宋_GB2312'

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ''
        expected_font = format_spec.get(style_name, {}).get('font_name')

        # 标题样式：清除 pandoc 或默认模板添加的蓝色
        is_heading = style_name.startswith('Heading ') if style_name else False

        for run in paragraph.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is None:
                continue

            # 标题 run：移除直接颜色设置（pandoc 默认蓝色）
            if is_heading:
                color_el = rPr.find(qn('w:color'))
                if color_el is not None:
                    rPr.remove(color_el)

            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                continue

            # 检查当前东亚字体
            current_east = rFonts.get(qn('w:eastAsia'))
            # 检查当前 ASCII 字体
            current_ascii = rFonts.get(qn('w:ascii'))

            if expected_font:
                # 有明确预期的字体，修正东亚字体
                if current_east and current_east != expected_font:
                    rFonts.set(qn('w:eastAsia'), expected_font)
                elif current_east is None:
                    rFonts.set(qn('w:eastAsia'), expected_font)
                # 修正 ASCII 字体（pandoc 有时会写入 'SimHei' 等）
                if current_ascii and expected_font == '仿宋_GB2312' and current_ascii != expected_font:
                    # 只在正文样式下修正 ASCII 字体
                    if style_name in ('Normal', 'Compact'):
                        rFonts.set(qn('w:ascii'), expected_font)

    # ── 图片段落行距修正 ──
    # pandoc 生成的内联图片承载于 Normal/Body Text 样式段落中，
    # 这些样式设置了固定值 28 磅行距（公文标准），
    # 但图片远高于 28 磅，会被裁剪并导致与后续文字重叠。
    # 解决方案：将包含图片的段落行距改为单倍行距（自动扩展以容纳图片）。
    _fix_image_paragraph_spacing(doc, format_spec)

    # ── 表格格式处理 ──
    if reference_docx_path:
        # 有模版：先复制段落格式（表格此时尚未插入，仅复制段落）
        doc.save(output_path)
        _copy_template_formatting(output_path, reference_docx_path,
                                  copy_paragraphs=True, copy_tables=False)
    else:
        # 无模版：应用公文标准表格格式
        for table in doc.tables:
            # ── 表格尺寸：根据窗口和内容自动调整 ──
            _set_table_autofit_window(table)

            for row in table.rows:
                for cell in row.cells:
                    # 设置单元格垂直居中
                    _set_cell_vertical_alignment(cell, 'center')

                    for paragraph in cell.paragraphs:
                        # 设置段落水平居中
                        paragraph.alignment = _get_alignment_enum('center')

                        # 清除首行缩进（防止从 Normal 样式继承 2 字符缩进）
                        pPr = paragraph._element.find(qn('w:pPr'))
                        if pPr is not None:
                            _clear_first_line_indent(pPr)

                        # 清理段落中所有 run 的文本多余空格
                        for run in paragraph.runs:
                            if run.text:
                                # 合并连续空白字符，去除首尾空白
                                cleaned = re.sub(r'\s+', ' ', run.text).strip()
                                run.text = cleaned

                        # 表格单元格内段落可能没有应用 Compact 样式
                        for run in paragraph.runs:
                            rPr = run._element.find(qn('w:rPr'))
                            if rPr is None:
                                continue

                            # 修正东亚字体
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                                rPr.insert(0, rFonts)
                            rFonts.set(qn('w:eastAsia'), DEFAULT_FONT)

                            # 修正字号为五号 (10.5pt)
                            sz = rPr.find(qn('w:sz'))
                            sz_cs = rPr.find(qn('w:szCs'))
                            target_sz = '21'  # 10.5pt = 21 half-pts
                            if sz is not None:
                                sz.set(qn('w:val'), target_sz)
                            if sz_cs is not None:
                                sz_cs.set(qn('w:val'), target_sz)

        doc.save(output_path)


# ──────────────────────────────────────────────
# 核心转换
# ──────────────────────────────────────────────

def _get_template_cache_dir():
    """获取参考模板缓存目录。"""
    cache_dir = os.path.join(str(Path.home()), '.cache', 'claude', 'md2docx')
    os.makedirs(cache_dir, exist_ok=True)
    return cache_dir


def convert_md_to_docx(
    input_path: str,
    output_path: str | None = None,
    reference_docx: str | None = None,
    extract_media: str | None = None,
    from_format: str = "markdown+pipe_tables+fenced_divs+bracketed_spans+raw_html",
    standalone: bool = True,
    extra_args: list[str] | None = None,
    skip_format: bool = False,
    no_cache_template: bool = False,
    format_spec_md: str | None = None,
):
    """
    使用 pandoc 将 Markdown 文件转换为 Word (.docx) 文件。

    默认会应用内置的公文格式规范（黑体标题、仿宋正文、固定行距28磅等），
    可通过 skip_format=True 跳过。

    Args:
        input_path: 输入的 Markdown 文件路径
        output_path: 输出的 .docx 文件路径（None 则自动生成）
        reference_docx: 额外的参考样式模板（叠加在内置格式之上）
        extract_media: 提取媒体文件的目录
        from_format: pandoc 输入格式
        standalone: 是否生成独立文档
        extra_args: 额外的 pandoc 参数
        skip_format: 跳过内置格式规范
        no_cache_template: 强制重新生成格式模板
        format_spec_md: 格式规范文件路径（如 格式要求.md），未指定则使用内置默认格式
    """
    # 加载格式规范
    format_spec = get_format_spec(format_spec_md)
    if not os.path.exists(input_path):
        print(f"错误：输入文件不存在: {input_path}")
        sys.exit(1)

    # 自动生成输出路径
    if output_path is None:
        input_stem = Path(input_path).stem
        input_dir = Path(input_path).parent
        output_path = str(input_dir / f"{input_stem}.docx")

    # ── HTML 表格检测与预处理 ──
    # 在 pandoc 转换前提取 HTML 表格，替换为占位符
    # (pandoc 会剥离 HTML 标签，无法保留表格及 colspan/rowspan 结构)
    original_md = Path(input_path).read_text(encoding='utf-8')
    modified_md, html_tables = extract_html_tables(original_md)

    if html_tables:
        temp_md = Path(input_path).with_suffix('.tmp.md')
        temp_md.write_text(modified_md, encoding='utf-8')
        actual_input = str(temp_md)
        print(f"HTML表格预处理: 检测到 {len(html_tables)} 个 HTML 表格，已替换为占位符")
    else:
        actual_input = input_path
        temp_md = None

    pandoc_bin = find_pandoc()
    cmd = [
        pandoc_bin,
        actual_input,
        "-o", output_path,
        "--from", from_format,
        # 将输入文件所在目录加入资源搜索路径，
        # 使 Markdown 中的相对路径图片引用能被 pandoc 正确解析
        "--resource-path", str(Path(actual_input).parent),
    ]

    if standalone:
        cmd.append("--standalone")

    # ── 内置格式模板 ──
    use_builtin_ref = not skip_format
    if use_builtin_ref:
        cache_dir = _get_template_cache_dir()
        builtin_ref_path = os.path.join(cache_dir, 'reference-template.docx')

        if no_cache_template or not os.path.exists(builtin_ref_path):
            print(f"生成内置格式模板: {builtin_ref_path}")
            create_reference_docx(builtin_ref_path, format_spec)
        else:
            print(f"使用缓存的格式模板: {builtin_ref_path}")

        # pandoc 允许多个 --reference-doc，后面的会覆盖前面的
        # 先加载内置模板，如果用户提供了额外模板则叠加
        cmd.extend(["--reference-doc", builtin_ref_path])

    # ── 用户提供的参考模板（叠加在内置之上） ──
    if reference_docx and os.path.exists(reference_docx):
        cmd.extend(["--reference-doc", reference_docx])

    if extract_media:
        cmd.extend(["--extract-media", extract_media])

    if extra_args:
        cmd.extend(extra_args)

    print(f"执行命令: {' '.join(cmd)}")
    print(f"输入文件: {input_path}")
    print(f"输出文件: {output_path}")

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            print(f"转换失败 (退出码 {result.returncode}):")
            print(result.stderr)
            sys.exit(1)
        if result.stderr:
            print(f"pandoc 信息:\n{result.stderr}")
    except FileNotFoundError:
        print("错误：未找到 pandoc。请确保 pandoc 已安装并在 PATH 中。")
        print("安装方法: https://pandoc.org/installing.html")
        sys.exit(1)
    except subprocess.TimeoutExpired:
        print("错误：转换超时（120秒）")
        sys.exit(1)

    # ── 后处理：修正东亚字体 ──
    if not skip_format:
        print("后处理：修正东亚字体与表格格式...")
        post_process_docx(
            output_path,
            format_spec,
            reference_docx_path=reference_docx if reference_docx else None,
        )

    # ── 后处理：插入 HTML 表格 ──
    if html_tables:
        _insert_html_tables(output_path, html_tables, format_spec)

    # ── 后处理：从模版复制表格格式（必须在 HTML 表格插入之后） ──
    if not skip_format and reference_docx and os.path.exists(reference_docx):
        _copy_template_formatting(
            output_path,
            reference_docx,
            copy_paragraphs=False,
            copy_tables=True,
        )

    # 清理临时文件
    if temp_md is not None:
        try:
            os.unlink(temp_md)
        except OSError:
            pass

    print(f"✅ 转换成功: {output_path}")
    return output_path


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────

def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="使用 pandoc 将 Markdown 转换为 Word (.docx)，输出符合公文格式规范",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    # 基本转换（使用内置默认格式）
    python scripts/convert_md2docx.py report.md

    # 从 格式要求.md 动态加载格式规范
    python scripts/convert_md2docx.py report.md -f 格式要求.md

    # 指定输出路径
    python scripts/convert_md2docx.py report.md -o output/report.docx

    # 格式规范 + 额外参考模板叠加
    python scripts/convert_md2docx.py report.md -f 格式要求.md -r custom-reference.docx

    # 退回纯 pandoc 行为（不应用格式规范）
    python scripts/convert_md2docx.py report.md --no-format

格式规范说明：
    未指定 --format-spec 时，使用内置默认格式（黑体标题、仿宋正文、固定行距28磅）。
    指定 --format-spec 后，从格式要求文件中动态加载样式定义。
        """,
    )
    parser.add_argument("input", help="输入的 Markdown 文件路径")
    parser.add_argument("-o", "--output", default=None, help="输出 .docx 文件路径")
    parser.add_argument("-f", "--format-spec", default=None,
                        help="格式规范文件路径（如 格式要求.md），未指定则使用内置默认格式")
    parser.add_argument("-r", "--reference", default=None,
                        help="额外的参考样式模板 .docx（叠加在内置格式之上）")
    parser.add_argument("--extract-media", default=None, help="提取媒体文件的目录")
    parser.add_argument("--no-standalone", action="store_true", help="不生成独立文档")
    parser.add_argument("--from", dest="from_format", default=None,
                        help="输入格式（默认自动检测）")
    parser.add_argument("--extra-args", nargs="*", default=None,
                        help="传递给 pandoc 的额外参数")
    parser.add_argument("--no-format", action="store_true",
                        help="跳过内置格式规范（退回纯 pandoc 行为）")
    parser.add_argument("--no-cache-template", action="store_true",
                        help="强制重新生成内置格式模板")

    args = parser.parse_args()
    extra = args.extra_args or []

    if not args.from_format:
        args.from_format = "markdown+pipe_tables+fenced_divs+bracketed_spans+raw_html"

    convert_md_to_docx(
        input_path=args.input,
        output_path=args.output,
        reference_docx=args.reference,
        extract_media=args.extract_media,
        from_format=args.from_format,
        standalone=not args.no_standalone,
        extra_args=extra,
        skip_format=args.no_format,
        no_cache_template=args.no_cache_template,
        format_spec_md=args.format_spec,
    )


if __name__ == "__main__":
    main()
